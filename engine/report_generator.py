import requests
import datetime
import io

class ReportGenerator:
    def __init__(self, ai_engine):
        self.ai = ai_engine
        
    def generate_weekly_report(self, assets_data: list, lang: str = "it") -> str:
        """
        Generates a professional financial report based on a list of asset data.
        """
        date_str = datetime.date.today().strftime("%d %B %Y")
        language_ctx = "Italian" if lang == "it" else "English"
        
        data_summary = ""
        for asset in assets_data:
            data_summary += f"- {asset['ticker']}: Price ${asset['price']:,.2f}, RSI {asset['rsi']:.1f}, Vol {asset['vol']:.1%}, Sentiment {asset['sentiment']}\n"
            
        system_prompt = f"""
        System: You are 'FinanceSensei', a top-tier institutional financial strategist.
        Task: Generate a 'Weekly Alpha Recap' report.
        Tone: Deeply professional, objective, and institutional.
        Language: ALWAYS respond in {language_ctx}.
        Structure:
        1. Macro Outlook: A 2-paragraph synthesis of the current market regime.
        2. Top Alpha Signals: Highlight the 2 most interesting assets from the data provided.
        3. Risk Assessment: Briefly mention thematic risks (liquidity, volatility).
        4. Conclusion: A final closing strategic thought.
        """
        
        user_prompt = f"Data Summary for this week ({date_str}):\n{data_summary}\n\nGenerate the complete report in Markdown."
        
        try:
            if self.ai.provider == "ollama":
                response = requests.post(self.ai.ollama_url, 
                                         json={"model": self.ai.model, "prompt": f"{system_prompt}\nUser: {user_prompt}", "stream": False},
                                         timeout=30)
                if response.status_code == 200:
                    report_content = response.json().get('response', "Failed to generate report content.")
                else:
                    report_content = "Connection to AI Engine lost during synthesis."
            else:
                report_content = self._heuristic_report(assets_data, lang)
        except Exception as e:
            report_content = f"Error during report generation: {str(e)}"
            
        return f"# FinanceSensei Weekly Alpha Report\n**Date: {date_str}**\n\n---\n\n{report_content}"

    def generate_docx_report(self, markdown_text: str) -> io.BytesIO:
        """
        Converts basic markdown report to a professional Word document.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is not installed. Run 'pip install python-docx'.")

        doc = Document()
        
        # Split by sections (simple parsing)
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                h = doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], level=2)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _heuristic_report(self, assets_data, lang):
        if lang == "it":
            return "Note: L'intelligenza profonda (Ollama) è necessaria per la sintesi dei report settimanali. Ecco un riassunto dei dati:\n" + \
                   "\n".join([f"- {a['ticker']}: RSI {a['rsi']:.1f}" for a in assets_data])
        return "Note: Deep Intelligence (Ollama) is required for strategic report synthesis. Data summary:\n" + \
               "\n".join([f"- {a['ticker']}: RSI {a['rsi']:.1f}" for a in assets_data])
